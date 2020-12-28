# -*- coding: utf-8 -*-
__docformat__ = "restructuredtext en"

import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.document import Document as _Document
from docx.document import Section as _Section
from docx.text.paragraph import Paragraph

"""
replace text in office docs
File:           ms_text_replace.py
Copyright:      Blink AG   
Author:         Steffen Kube <steffen@blink-dx.com>
"""

class MS_sub_object:
    
    def __init__(self, docobj):
        '''
        ... can be a sub object in a document
        '''
        self.doc = docobj
        self.warnings=[]
        
    def _add_warning(self, text):
        self.warnings.append(text)
        
    def get_text_raw(self):
        '''
        get text from document: 
        - tables in the right order
        '''

        def iter_block_items(parent):
            """
            Generate a reference to each paragraph and table child within *parent*,
            in document order. Each returned value is an instance of either Table or
            Paragraph. *parent* would most commonly be a reference to a main
            Document object, but also works for a _Cell object, which itself can
            contain paragraphs and tables.
            """
            typex='STD'
            if isinstance(parent, _Document):
                parent_elm = parent.element.body
            if isinstance(parent, docx.section._Header):
                parent_elm = parent
                typex='Header'
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            elif isinstance(parent, _Row):
                parent_elm = parent._tr
            else:
                raise ValueError("something's not right")
            
            if typex=='Header':
                
                for child in parent_elm.tables:
                    yield child
                    
            else:
                
                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)
                    else:
                        a=1
                    
        document = self.doc
        text=''
        
        for block in iter_block_items(document):
        
            #print(block.text if isinstance(block, Paragraph) else '<table>')
            if isinstance(block, Paragraph):
                text = text + block.text + "\n" 
                
            elif isinstance(block, Table):
                
                for row in block.rows:
                    
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = text + paragraph.text + "\t" 
                    text = text + "\n" 

        return text
    
    def row_replace_one_key(self, old, new):
        '''
        xobj must be a row
        :return: int found : 0,1
        '''
        found=0
        xobj = self.doc 
        
        for row in xobj.rows:

            for cell in row.cells:
                cell_proc = cell.text.strip()
                if cell_proc==old:
                    cell_new  = cell_proc.replace(old,new)
                    cell.text = cell_new
                    found=1
                    break
            if found:
                break
        
        if not found:
            self._add_warning('Key "'+old+'" not found in document (row_replace)')
            
        return found
     
    
    def table_find_replace_many(self, replace_list):
        '''
        update row cells
        :param replace_list: list of {'old':, 'new':}
        '''  
        
        for rep_row in replace_list:
            found = self.row_replace_one_key(rep_row['old'],  rep_row['new']) 
    
    def get_warnings(self):
        return self.warnings    

class MS_text_replace:
    
    def __init__(self, filename):
        self.filename = filename
        self.doc = docx.Document(self.filename)
        
    def get_header(self):
        header = self.doc.sections[0].header
        return header
    
        
    def update_cells(self, xobj, replace_list):
        '''
        update row cells
        :param replace_list: list of {'old':, 'new':}
        '''        
        i=0
        for cell in xobj.cells:
            if 'new' in replace_list[i]:
                cell.text = replace_list[i]['new']
            i=i+1
            
    
    
    def table_find_row(self, pos_key, row_cnt):
        '''
        :param pos_key: position key, to identify the right row
        :param row_cnt: if occurs many time, find the row_cnt match (starts at 0)
        :return {} table_row_COORD or None
          'table_cnt':table_cnt, 
          'row_cnt':row_cnt
          'content_row_poi' : content pointer
        '''
        
        # tables 
        found = 0
        table_cnt= 0
        tables   = self.doc.tables 
        
        for table in tables:
            
            row_cnt=0
            for i, row in enumerate(table.rows):

                for cell in row.cells:
                    cell_proc=cell.text.strip()
                    if cell_proc==pos_key:
                        found = 1
                        break
                    
                if found:
                    break
                row_cnt = row_cnt + 1
            if found:
                break            
            table_cnt = table_cnt + 1
        
        if not found:
            return None
        
        return {'table_cnt':table_cnt, 'row_cnt':row_cnt, 'content_row_poi': row }             
    
    def table_add_row(self, table_row_COORD):
        '''
        :param table_row_COORD: {} position info, to identify the position where the new row comes AFTER 
        :param object content_obj: row object
        :return: row_poi of new row
        '''
          
        found   = 0
        row_poi = 0
        table_cnt= 0
        tables   = self.doc.tables 
    
        for table in tables:
            
            if table_row_COORD['table_cnt']==table_cnt:
                
                row_cnt=0
                for i, row in enumerate(table.rows):
        
                    if table_row_COORD['row_cnt']==row_cnt:
                        found=1
                        row_poi = table.add_row()
                        break
                    row_cnt = row_cnt + 1
                    
            if found:
                break
            table_cnt = table_cnt + 1 
            
        return row_poi
    
    def get_header_table(self):
        '''
        get table obj
        '''
        header = self.doc.sections[0].header
        tables =  header.tables
        for table_obj in tables:
            break
    
        return table_obj
    
    def save(self, filename=''):
        if filename=='':
            filename=self.filename
        self.doc.save(filename)
    
if __name__ == "__main__":
    
    filename    =r'C:\Users\Steffen\Documents\Code\blinkdms\test\blinkdms\code\lib\oVERSION\Vorlage_SOP_t1.docx'
    filename_new=r'C:\Users\Steffen\Documents\Code\blinkdms\test\blinkdms\code\lib\oVERSION\Vorlage_SOP_t1.tr.docx'
    version_id=0
    
    a = MS_text_replace(filename)
    header = a.get_header()
    doc_sub_obj_lib = MS_sub_object(header)
    text1 = doc_sub_obj_lib.get_text_raw()
    warnings =  doc_sub_obj_lib.get_warnings()
    
    replace_list = [
        { 'old': '{{SOP-Titel}}', 'new': 'Hello Du titl' },
        { 'old': '{{SOP-Code}}', 'new': 'SOP-0404-wwww' },
        { 'old': '{{SOP-Version}}', 'new': '3.0' },
        { 'old': '{{SOP-Valid}}', 'new': '2020-08-11' },
    ]      
    
    header_tab = a.get_header_table()
    doc_sub2_obj_lib = MS_sub_object(header_tab)
    doc_sub2_obj_lib.table_find_replace_many(replace_list)
    
    a.save(filename_new)
    
    print('READY')